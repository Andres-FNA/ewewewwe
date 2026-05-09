"""
document_loader.py
Carga documentos de texto plano (.txt), PDF (.pdf) y Word (.docx),
y los divide en chunks con solapamiento.

Dependencias:
  pip install pypdf python-docx
"""

import os
import re
from typing import List, Dict


# ─────────────────────────────────────────────
# Tipos de datos
# ─────────────────────────────────────────────

class Chunk:
    """Representa un fragmento de texto con metadatos."""
    def __init__(self, text: str, source: str, chunk_id: int):
        self.text     = text
        self.source   = source      # nombre del archivo original
        self.chunk_id = chunk_id

    def __repr__(self):
        preview = self.text[:60].replace("\n", " ")
        return f"Chunk(id={self.chunk_id}, source='{self.source}', text='{preview}...')"


# ─────────────────────────────────────────────
# Lectores por tipo de archivo
# ─────────────────────────────────────────────

def load_txt(path: str) -> str:
    """Lee un archivo .txt y devuelve su contenido."""
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def load_pdf(path: str) -> str:
    """Extrae texto de un PDF usando pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("pypdf no está instalado. Instálalo con: pip install pypdf")

    reader = PdfReader(path)
    pages_text = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages_text.append(text)

    full_text = "\n\n".join(pages_text)

    if not full_text.strip():
        print(f"  [AVISO] '{os.path.basename(path)}': sin texto extraíble (¿PDF escaneado?). Se omite.")
        return ""

    return full_text


def load_docx(path: str) -> str:
    """Extrae texto de un archivo Word (.docx) usando python-docx."""
    try:
        import docx
    except ImportError:
        raise ImportError("python-docx no está instalado. Instálalo con: pip install python-docx")

    doc   = docx.Document(path)
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                parts.append(" | ".join(row_texts))

    return "\n\n".join(parts)


# ─────────────────────────────────────────────
# Carga de documentos desde carpeta
# ─────────────────────────────────────────────

SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}


def load_documents_from_folder(folder: str) -> List[Dict]:
    """
    Carga todos los archivos soportados (.txt, .pdf, .docx) de una carpeta.
    Devuelve lista de dicts con keys: 'source', 'text'.
    """
    documents = []

    if not os.path.isdir(folder):
        print(f"  [ERROR] La carpeta '{folder}' no existe.")
        return documents

    for filename in sorted(os.listdir(folder)):
        ext = os.path.splitext(filename)[1].lower()
        if ext not in SUPPORTED_EXTENSIONS:
            continue

        filepath = os.path.join(folder, filename)

        try:
            if ext == ".txt":
                text = load_txt(filepath)
            elif ext == ".pdf":
                text = load_pdf(filepath)
            elif ext == ".docx":
                text = load_docx(filepath)
        except Exception as e:
            print(f"  [ERROR] No se pudo leer '{filename}': {e}")
            continue

        if not text.strip():
            print(f"  [OMITIDO] '{filename}' — sin contenido útil")
            continue

        documents.append({"source": filename, "text": text})
        print(f"  [CARGADO] {filename} ({ext.upper()[1:]}) — {len(text)} caracteres")

    return documents


# ─────────────────────────────────────────────
# Limpieza y chunking
# ─────────────────────────────────────────────
def clean_text(text: str) -> str:
    """
    Limpieza avanzada para PDFs académicos.

    Mejora:
    - elimina saltos raros de PDF
    - evita palabras pegadas incorrectamente
    - elimina espacios múltiples
    - elimina líneas basura comunes
    - mejora chunks para retrieval + RAGAS
    """

    if not text:
        return ""

    # Normalizar saltos de línea
    text = text.replace("\r\n", "\n")
    text = text.replace("\r", "\n")

    # Eliminar encabezados comunes del PDF
    basura_pdf = [
        "FR-PD-G-501",
        "Programa de Asignatura",
        "Versión 6.0",
        "Optimización y Mejoramiento",
        "Página 1 de 5",
        "Página 2 de 5",
        "Página 3 de 5",
        "Página 4 de 5",
        "Página 5 de 5",
    ]

    for basura in basura_pdf:
        text = text.replace(basura, "")

    # Corregir espacios múltiples
    text = re.sub(r"[ \t]+", " ", text)

    # Corregir demasiados saltos
    text = re.sub(r"\n{3,}", "\n\n", text)

    # IMPORTANTE:
    # evitar que palabras de líneas distintas
    # se peguen mal como:
    #
    # Análisis y Diseño de Sistemas
    # Arquitectura
    #
    # Aquí intentamos mantener separación lógica
    text = re.sub(
        r"([a-záéíóúñ])\n([A-ZÁÉÍÓÚÑ])",
        r"\1. \2",
        text
    )

    # Eliminar espacios al inicio/final
    text = text.strip()

    return text


def chunk_text(
    text: str,
    source: str,
    chunk_size: int = 500,
    overlap: int = 100,
    start_id: int = 0,
) -> List[Chunk]:
    """
    Divide texto en chunks de `chunk_size` caracteres con `overlap` de solapamiento.
    El solapamiento preserva contexto entre chunks consecutivos.
    """
    text     = clean_text(text)
    chunks   = []
    pos      = 0
    chunk_id = start_id

    while pos < len(text):
        end      = pos + chunk_size
        fragment = text[pos:end]

        # Intentar cortar en un espacio o salto de línea natural
        if end < len(text):
            last_break = max(fragment.rfind(" "), fragment.rfind("\n"))
            if last_break > chunk_size // 2:
                fragment = fragment[:last_break]
                end      = pos + last_break

        chunks.append(Chunk(
            text     = fragment.strip(),
            source   = source,
            chunk_id = chunk_id,
        ))
        chunk_id += 1
        pos       = end - overlap   # retroceso = solapamiento

    return chunks


def build_chunks_from_documents(
    documents: List[Dict],
    chunk_size: int = 500,
    overlap: int = 100,
) -> List[Chunk]:
    """
    Procesa una lista de documentos y devuelve todos los chunks.
    """
    all_chunks = []
    for doc in documents:
        doc_chunks = chunk_text(
            text       = doc["text"],
            source     = doc["source"],
            chunk_size = chunk_size,
            overlap    = overlap,
            start_id   = len(all_chunks),
        )
        all_chunks.extend(doc_chunks)
        print(f"  [CHUNKED] {doc['source']} → {len(doc_chunks)} chunks")
    return all_chunks